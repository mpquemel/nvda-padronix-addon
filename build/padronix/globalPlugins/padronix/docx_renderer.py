# -*- coding: utf-8 -*-
# Padronix - NVDA Add-on for Academic Paper Formatting
# Copyright (C) 2026 Melhym Pereira Quemel
# Email: mpquemel@gmail.com

"""Renderizador de documentos DOCX"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import addonHandler
addonHandler.initTranslation()

class DocxRenderer:
    """Renderiza estrutura JSON em documento DOCX formatado"""
    
    def __init__(self, standard='abnt'):
        self.standard = standard
        self.styles = self._load_styles(standard)
    
    def _load_styles(self, standard):
        """Carregar estilos do padrão acadêmico"""
        styles_dir = os.path.join(os.path.dirname(__file__), "..", "..", "styles")
        styles_file = os.path.join(styles_dir, f"{standard}.json")
        
        if os.path.exists(styles_file):
            with open(styles_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            # Estilos padrão ABNT
            return {
                "font": "Arial",
                "font_size": 12,
                "line_spacing": 1.5,
                "margins": {
                    "top": 3,
                    "bottom": 2,
                    "left": 3,
                    "right": 2
                }
            }
    
    def render(self, structured_data):
        """Renderizar documento completo"""
        
        doc = Document()
        
        # Configurar margens
        section = doc.sections[0]
        section.top_margin = Cm(self.styles['margins']['top'])
        section.bottom_margin = Cm(self.styles['margins']['bottom'])
        section.left_margin = Cm(self.styles['margins']['left'])
        section.right_margin = Cm(self.styles['margins']['right'])
        
        # Título
        if 'title' in structured_data:
            title = doc.add_heading(structured_data['title'], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resumo
        if 'abstract' in structured_data:
            doc.add_heading('Resumo', level=2)
            doc.add_paragraph(structured_data['abstract'])
        
        # Palavras-chave
        if 'keywords' in structured_data:
            keywords = ', '.join(structured_data['keywords'])
            doc.add_paragraph(f'Palavras-chave: {keywords}')
        
        # Seções
        if 'sections' in structured_data:
            for section_data in structured_data['sections']:
                heading_level = section_data.get('level', 1)
                heading = doc.add_heading(
                    section_data.get('heading', ''),
                    level=min(heading_level, 9)
                )
                
                if 'content' in section_data:
                    doc.add_paragraph(section_data['content'])
        
        # Referências
        if 'references' in structured_data:
            doc.add_heading('Referências', level=1)
            for ref in structured_data['references']:
                doc.add_paragraph(ref, style='List Bullet')
        
        # Salvar documento
        output_dir = os.path.join(os.environ.get('TEMP', 'C:\\Temp'))
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Padronix_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)
        
        doc.save(output_path)
        
        return output_path
