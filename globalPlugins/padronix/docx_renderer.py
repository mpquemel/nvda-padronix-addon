# -*- coding: utf-8 -*-
"""
Padronix - Renderizador de DOCX
Copyright (C) 2026 Melhym Pereira Quemel <mpquemel@gmail.com>

Renderiza dados estruturados em documento DOCX formatado.
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import addonHandler

addonHandler.initTranslation()

import logHandler


class DocxRenderer:
    """
    Renderiza estrutura JSON em documento DOCX formatado.
    
    Suporta múltiplos padrões acadêmicos com configurações dinâmicas.
    """
    
    def __init__(self, standard: str = "abnt") -> None:
        """
        Inicializa renderizador.
        
        Args:
            standard: Padrão acadêmico (abnt, apa, ieee, etc.)
        """
        self.standard = standard
        self.styles = self._load_styles(standard)
    
    def _load_styles(self, standard: str) -> Dict[str, Any]:
        """
        Carrega estilos do padrão acadêmico.
        
        Args:
            standard: Nome do padrão
            
        Returns:
            Dicionário com configurações de estilo
        """
        # Caminho para arquivo de estilos
        plugin_dir = Path(__file__).parent
        styles_file = plugin_dir / ".." / ".." / "styles" / f"{standard}.json"
        
        if styles_file.exists():
            try:
                with open(styles_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logHandler.log.warning(f"Padronix: Erro ao carregar estilo {standard}: {e}")
        
        # Fallback para estilos padrão ABNT
        return self._get_default_styles()
    
    def _get_default_styles(self) -> Dict[str, Any]:
        """Retorna estilos padrão ABNT."""
        return {
            "font": "Arial",
            "font_size": 12,
            "line_spacing": 1.5,
            "margins": {
                "top": 3.0,
                "bottom": 2.0,
                "left": 3.0,
                "right": 2.0
            },
            "title": {
                "font_size": 14,
                "bold": True,
                "alignment": "center"
            },
            "heading1": {
                "font_size": 14,
                "bold": True
            },
            "heading2": {
                "font_size": 12,
                "bold": True
            },
            "paragraph": {
                "alignment": "justify",
                "first_line_indent": 1.25  # cm
            }
        }
    
    def render(self, structured_data: Dict[str, Any]) -> str:
        """
        Renderiza documento completo.
        
        Args:
            structured_data: Dicionário com dados estruturados do artigo
            
        Returns:
            Caminho completo do arquivo DOCX gerado
        """
        try:
            from docx import Document
            from docx.shared import Cm, Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document()
            
            # Configurar margens
            section = doc.sections[0]
            section.top_margin = Cm(self.styles['margins']['top'])
            section.bottom_margin = Cm(self.styles['margins']['bottom'])
            section.left_margin = Cm(self.styles['margins']['left'])
            section.right_margin = Cm(self.styles['margins']['right'])
            
            # Configurar fonte padrão
            style = doc.styles['Normal']
            font = style.font
            font.name = self.styles['font']
            font.size = Pt(self.styles['font_size'])
            
            # === TÍTULO ===
            if 'title' in structured_data and structured_data['title']:
                title = doc.add_heading(structured_data['title'], level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_format = title.runs[0]
                title_format.font.size = Pt(self.styles['title']['font_size'])
                title_format.font.bold = True
                doc.add_paragraph()  # Espaço
            
            # === RESUMO ===
            if 'abstract' in structured_data and structured_data['abstract']:
                doc.add_heading(_('Resumo'), level=1)
                abstract_para = doc.add_paragraph(structured_data['abstract'])
                abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # === ABSTRACT (inglês) ===
            if 'abstract_en' in structured_data and structured_data['abstract_en']:
                doc.add_heading('Abstract', level=1)
                abstract_en_para = doc.add_paragraph(structured_data['abstract_en'])
                abstract_en_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # === PALAVRAS-CHAVE ===
            if 'keywords' in structured_data and structured_data['keywords']:
                keywords = structured_data['keywords']
                if isinstance(keywords, list):
                    keywords_text = ', '.join(keywords)
                else:
                    keywords_text = str(keywords)
                
                doc.add_paragraph()
                kw_para = doc.add_paragraph()
                kw_para.add_run(_('Palavras-chave: ')).bold = True
                kw_para.add_run(keywords_text)
            
            # === KEYWORDS (inglês) ===
            if 'keywords_en' in structured_data and structured_data['keywords_en']:
                keywords_en = structured_data['keywords_en']
                if isinstance(keywords_en, list):
                    keywords_en_text = ', '.join(keywords_en)
                else:
                    keywords_en_text = str(keywords_en)
                
                doc.add_paragraph()
                kw_en_para = doc.add_paragraph()
                kw_en_para.add_run('Keywords: ').bold = True
                kw_en_para.add_run(keywords_en_text)
            
            doc.add_page_break()
            
            # === SEÇÕES ===
            if 'sections' in structured_data and structured_data['sections']:
                for section_data in structured_data['sections']:
                    self._render_section(doc, section_data)
            
            # === REFERÊNCIAS ===
            if 'references' in structured_data and structured_data['references']:
                doc.add_page_break()
                doc.add_heading(_('Referências'), level=1)
                
                references = structured_data['references']
                if isinstance(references, list):
                    for ref in references:
                        ref_para = doc.add_paragraph(ref, style='List Bullet')
                        ref_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    doc.add_paragraph(str(references))
            
            # === SALVAR DOCUMENTO ===
            output_path = self._generate_output_path()
            doc.save(output_path)
            
            logHandler.log.info(f"Padronix: Documento gerado em {output_path}")
            return output_path
            
        except ImportError as e:
            logHandler.log.error(f"Padronix: python-docx não encontrado: {e}")
            raise RuntimeError(_("Biblioteca python-docx não instalada"))
            
        except Exception as e:
            logHandler.log.error(f"Padronix: Erro ao renderizar DOCX: {e}")
            raise
    
    def _render_section(self, doc: Any, section_data: Dict[str, Any]) -> None:
        """
        Renderiza uma seção do artigo.
        
        Args:
            doc: Documento DOCX
            section_data: Dados da seção
        """
        heading = section_data.get('heading', '')
        level = section_data.get('level', 1)
        content = section_data.get('content', '')
        
        # Adicionar cabeçalho
        heading_level = min(level + 1, 9)  # DOCX suporta até heading 9
        doc.add_heading(heading, level=heading_level)
        
        # Adicionar conteúdo
        if content:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            para = doc.add_paragraph(content)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Indentação da primeira linha (ABNT)
            if self.styles.get('paragraph', {}).get('first_line_indent'):
                from docx.shared import Cm
                para.paragraph_format.first_line_indent = Cm(
                    self.styles['paragraph']['first_line_indent']
                )
    
    def _generate_output_path(self) -> str:
        """
        Gera caminho de saída único para o documento.
        
        Returns:
            Caminho completo do arquivo
        """
        # Diretório de saída (Desktop ou Temp)
        desktop = Path.home() / "Desktop"
        if not desktop.exists():
            desktop = Path(os.environ.get('TEMP', 'C:\\Temp'))
        
        # Timestamp único
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Padronix_{timestamp}.docx"
        
        output_path = desktop / filename
        return str(output_path)
